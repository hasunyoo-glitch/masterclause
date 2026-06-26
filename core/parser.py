"""Document parsing: PDF / DOCX → normalized text (+ light structure).

Primary PDF path is PyMuPDF (fast, layout-aware); pdfplumber is a fallback.
DOCX uses python-docx. The goal is faithful, verbatim text the analyzer can
quote exactly — so we normalize whitespace conservatively and never drop
content.
"""
from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path

SUPPORTED_SUFFIXES = {".pdf", ".docx"}


class ParserError(Exception):
    """Raised when a document cannot be read or is an unsupported type."""


@dataclass
class ParsedDocument:
    path: str
    file_format: str          # "pdf" | "docx"
    text: str                 # normalized full text (verbatim content preserved)
    page_count: int = 0       # PDF pages (0 for docx)
    paragraph_count: int = 0
    char_count: int = field(default=0)

    def preview(self, limit: int = 600) -> str:
        snippet = self.text.strip()
        return snippet[:limit] + ("…" if len(snippet) > limit else "")


# --------------------------------------------------------------------------- #
# Public API
# --------------------------------------------------------------------------- #
def parse(path: str | Path) -> ParsedDocument:
    p = Path(path)
    if not p.exists():
        raise ParserError(f"파일을 찾을 수 없습니다: {p}")
    suffix = p.suffix.lower()
    if suffix not in SUPPORTED_SUFFIXES:
        raise ParserError(f"지원하지 않는 형식입니다: {suffix} (PDF/DOCX만 지원)")

    if suffix == ".pdf":
        text, pages = _parse_pdf(p)
        fmt = "pdf"
    else:
        text, pages = _parse_docx(p), 0
        fmt = "docx"

    text = _normalize(text)
    if not text.strip():
        raise ParserError(
            "문서에서 텍스트를 추출하지 못했습니다. "
            "스캔 이미지 PDF라면 OCR이 필요합니다."
        )
    return ParsedDocument(
        path=str(p),
        file_format=fmt,
        text=text,
        page_count=pages,
        paragraph_count=_count_paragraphs(text),
        char_count=len(text),
    )


# --------------------------------------------------------------------------- #
# PDF
# --------------------------------------------------------------------------- #
def _parse_pdf(p: Path) -> tuple[str, int]:
    text, pages = _parse_pdf_pymupdf(p)
    if text.strip():
        return text, pages
    # Fallback for PDFs PyMuPDF struggles with.
    return _parse_pdf_pdfplumber(p)


def _parse_pdf_pymupdf(p: Path) -> tuple[str, int]:
    try:
        import fitz  # PyMuPDF
    except ImportError:
        return "", 0
    try:
        parts: list[str] = []
        with fitz.open(p) as doc:
            page_count = doc.page_count
            for page in doc:
                parts.append(page.get_text("text"))
        return "\n".join(parts), page_count
    except Exception as exc:  # pragma: no cover - defensive
        raise ParserError(f"PDF 파싱 실패(PyMuPDF): {exc}") from exc


def _parse_pdf_pdfplumber(p: Path) -> tuple[str, int]:
    try:
        import pdfplumber
    except ImportError as exc:
        raise ParserError(
            "PDF 파서를 사용할 수 없습니다. PyMuPDF 또는 pdfplumber를 설치하세요."
        ) from exc
    try:
        parts: list[str] = []
        with pdfplumber.open(p) as pdf:
            page_count = len(pdf.pages)
            for page in pdf.pages:
                parts.append(page.extract_text() or "")
        return "\n".join(parts), page_count
    except Exception as exc:  # pragma: no cover - defensive
        raise ParserError(f"PDF 파싱 실패(pdfplumber): {exc}") from exc


# --------------------------------------------------------------------------- #
# DOCX
# --------------------------------------------------------------------------- #
def _parse_docx(p: Path) -> str:
    try:
        import docx  # python-docx
    except ImportError as exc:
        raise ParserError("python-docx가 설치되지 않았습니다.") from exc
    try:
        document = docx.Document(str(p))
    except Exception as exc:  # pragma: no cover - defensive
        raise ParserError(f"DOCX 파싱 실패: {exc}") from exc

    blocks: list[str] = [para.text for para in document.paragraphs]
    # Include table cell text — contract terms often live in tables.
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                blocks.append(" | ".join(cells))
    return "\n".join(blocks)


# --------------------------------------------------------------------------- #
# Normalization helpers (conservative — preserve verbatim content)
# --------------------------------------------------------------------------- #
def _normalize(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    # Collapse runs of spaces/tabs but keep line structure for clause boundaries.
    text = re.sub(r"[ \t]+", " ", text)
    # Collapse 3+ blank lines to a single blank line.
    text = re.sub(r"\n{3,}", "\n\n", text)
    # Strip trailing spaces per line.
    text = "\n".join(line.rstrip() for line in text.split("\n"))
    return text.strip()


def _count_paragraphs(text: str) -> int:
    return sum(1 for block in re.split(r"\n\s*\n", text) if block.strip())
