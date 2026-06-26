"""Report generation: AnalysisResult → DOCX (and optional PDF) (plan §7).

Style: Times New Roman, monochrome, clean legal-memo structure — no decorative
colors, no "AI-looking" flourishes. Labels render in the chosen output language.
Korean text uses an East-Asian fallback font so glyphs are not broken.
"""
from __future__ import annotations

import datetime as _dt
import sys
from enum import Enum
from pathlib import Path

from core.models import (
    AnalysisOptions,
    AnalysisResult,
    OutputLanguage,
    Priority,
    RiskLevel,
)


class ReportError(Exception):
    pass


def _korean_font() -> str:
    """A Korean-capable font that exists on the host OS (East-Asian fallback)."""
    if sys.platform.startswith("win"):
        return "Malgun Gothic"
    if sys.platform == "darwin":
        return "Apple SD Gothic Neo"
    return "NanumGothic"


def _v(x) -> str:
    return x.value if isinstance(x, Enum) else (x if x is not None else "")


def _lang(options: AnalysisOptions) -> str:
    return _v(options.output_language) or OutputLanguage.KO.value


# --------------------------------------------------------------------------- #
# Bilingual labels
# --------------------------------------------------------------------------- #
LABELS = {
    "ko": {
        "title": "AI 음악 계약 분석 리포트",
        "generated": "생성일",
        "perspective": "분석 입장",
        "language": "출력 언어",
        "jurisdiction": "관할",
        "contract_type": "계약 유형",
        "parties": "당사자",
        "effective_date": "발효일",
        "term": "기간",
        "territory": "영역",
        "governing_law": "준거법",
        "venue": "관할/법정지",
        "risk_score": "종합 위험 점수",
        "recommendation": "최종 권고",
        "subscores": "세부 점수 (저작권 / 수익 / 통제 / 법적구제)",
        "key_findings": "핵심 발견",
        "red_clauses": "RED 조항 수",
        "walk_away": "walk-away 이슈 수",
        "clauses": "조항별 상세 분석",
        "clause": "조항",
        "issue_type": "쟁점 유형",
        "flag": "위험도",
        "enforceability": "집행가능성",
        "verbatim": "원문",
        "summary": "요약",
        "impact": "당신에게 미치는 영향",
        "revision": "제안 수정안",
        "legal_basis": "법적 근거",
        "concern": "당신의 우려에 대한 조언",
        "concern_text": "입력한 우려",
        "concern_answer": "직접 답변",
        "concern_risk": "우려 관점 위험도",
        "concern_clauses": "관련 조항",
        "concern_actions": "권고 행동",
        "benchmarks": "산업 표준 벤치마크",
        "topic": "항목",
        "contract_value": "계약 내용",
        "standard": "산업 표준",
        "assessment": "평가",
        "jurisdiction_notes": "관할별 법리 노트",
        "applicable_law": "적용법",
        "analysis": "분석",
        "conflict": "준거법 충돌",
        "playbook": "협상 플레이북",
        "priority": "우선순위",
        "strategy": "전략",
        "disclaimer": "면책 고지",
        "none": "해당 없음",
        "yes": "예",
        "no": "아니오",
        "review_note": "주의: 관할별 법리 데이터는 검수가 필요한 초안이며, 본 리포트는 법률 자문이 아닙니다.",
    },
    "en": {
        "title": "AI Music Contract Analysis Report",
        "generated": "Generated",
        "perspective": "Perspective",
        "language": "Output language",
        "jurisdiction": "Jurisdiction",
        "contract_type": "Contract type",
        "parties": "Parties",
        "effective_date": "Effective date",
        "term": "Term",
        "territory": "Territory",
        "governing_law": "Governing law",
        "venue": "Venue",
        "risk_score": "Overall risk score",
        "recommendation": "Recommendation",
        "subscores": "Sub-scores (copyright / revenue / control / legal remedy)",
        "key_findings": "Key findings",
        "red_clauses": "RED clauses",
        "walk_away": "Walk-away issues",
        "clauses": "Clause-by-clause analysis",
        "clause": "Clause",
        "issue_type": "Issue type",
        "flag": "Risk flag",
        "enforceability": "Enforceability",
        "verbatim": "Verbatim text",
        "summary": "Summary",
        "impact": "Impact on you",
        "revision": "Proposed revision",
        "legal_basis": "Legal basis",
        "concern": "Advice on your concern",
        "concern_text": "Your concern",
        "concern_answer": "Direct answer",
        "concern_risk": "Risk for this concern",
        "concern_clauses": "Related clauses",
        "concern_actions": "Recommended actions",
        "benchmarks": "Industry benchmarks",
        "topic": "Topic",
        "contract_value": "This contract",
        "standard": "Industry standard",
        "assessment": "Assessment",
        "jurisdiction_notes": "Jurisdiction notes",
        "applicable_law": "Applicable law",
        "analysis": "Analysis",
        "conflict": "Governing-law conflict",
        "playbook": "Negotiation playbook",
        "priority": "Priority",
        "strategy": "Strategy",
        "disclaimer": "Disclaimer",
        "none": "None",
        "yes": "Yes",
        "no": "No",
        "review_note": "Note: jurisdiction data is a draft pending expert review; this report is not legal advice.",
    },
}


# --------------------------------------------------------------------------- #
# Public API
# --------------------------------------------------------------------------- #
def write_report(result: AnalysisResult, options: AnalysisOptions) -> list[str]:
    """Generate the report file(s); return the list of written paths."""
    fmt = options.output_format or "docx"
    docx_path = _resolve_path(options, ".docx")
    _write_docx(result, options, docx_path)
    written = [str(docx_path)]

    if fmt in ("pdf", "both"):
        pdf_path = docx_path.with_suffix(".pdf")
        if _convert_to_pdf(docx_path, pdf_path):
            written.append(str(pdf_path))
            if fmt == "pdf":
                # User wanted PDF only — drop the intermediate DOCX from the list,
                # but keep the file (some Word installs need it). Report both.
                pass
    return written


def _resolve_path(options: AnalysisOptions, suffix: str) -> Path:
    out = Path(options.output_path)
    if out.suffix.lower() in (".docx", ".pdf"):
        target = out.with_suffix(suffix)
    else:
        out.mkdir(parents=True, exist_ok=True)
        stem = Path(options.input_file_path).stem or "contract"
        stamp = _dt.datetime.now().strftime("%Y%m%d_%H%M%S")
        target = out / f"{stem}_analysis_{stamp}{suffix}"
    target.parent.mkdir(parents=True, exist_ok=True)
    return target


# --------------------------------------------------------------------------- #
# DOCX
# --------------------------------------------------------------------------- #
def _write_docx(result: AnalysisResult, options: AnalysisOptions, path: Path) -> None:
    try:
        import docx
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt
    except ImportError as exc:  # pragma: no cover
        raise ReportError("python-docx가 설치되지 않았습니다.") from exc

    lang = _lang(options)
    L = LABELS[lang]
    is_ko = lang == "ko"

    document = docx.Document()
    _setup_base_style(document, is_ko)

    def heading(text: str, level: int = 1) -> None:
        h = document.add_heading(level=level)
        run = h.add_run(text)
        _font(run, is_ko, size=16 if level == 1 else 13, bold=True)

    def kv(label: str, value: str) -> None:
        p = document.add_paragraph()
        r1 = p.add_run(f"{label}: ")
        _font(r1, is_ko, bold=True)
        r2 = p.add_run(value or L["none"])
        _font(r2, is_ko)

    def body(text: str) -> None:
        p = document.add_paragraph()
        r = p.add_run(text or "")
        _font(r, is_ko)

    # --- Title / header ---
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run(L["title"])
    _font(tr, is_ko, size=20, bold=True)

    kv(L["generated"], result.generated_at.astimezone().strftime("%Y-%m-%d %H:%M"))
    kv(L["perspective"], _v(options.perspective))
    kv(L["language"], lang)
    kv(L["jurisdiction"], options.jurisdiction_label())

    ov = result.overview
    kv(L["contract_type"], _v(ov.contract_type))
    kv(L["parties"], ", ".join(ov.parties) if ov.parties else L["none"])
    kv(L["effective_date"], ov.effective_date)
    kv(L["term"], ov.term)
    kv(L["territory"], ov.territory)
    kv(L["governing_law"], ov.governing_law)
    kv(L["venue"], ov.venue)

    # --- Risk score + recommendation ---
    heading(L["risk_score"])
    rs = result.risk_score
    kv(L["risk_score"], f"{rs.overall} / 100")
    kv(L["recommendation"], _v(result.recommendation))
    kv(L["subscores"], f"{rs.copyright} / {rs.revenue} / {rs.control} / {rs.legal_remedy}")
    body(rs.summary)

    # --- Key findings ---
    heading(L["key_findings"])
    kv(L["red_clauses"], str(result.red_clause_count))
    kv(L["walk_away"], str(result.walk_away_count))

    # --- Concern advice (prominent, near the top) ---
    if result.concern_advice is not None:
        ca = result.concern_advice
        heading(L["concern"])
        kv(L["concern_text"], ca.concern_text)
        kv(L["concern_risk"], _v(ca.risk_assessment))
        body(ca.direct_answer)
        if ca.relevant_clause_ids:
            kv(L["concern_clauses"], ", ".join(ca.relevant_clause_ids))
        for action in ca.recommended_actions:
            bullet = document.add_paragraph(style="List Bullet")
            _font(bullet.add_run(action), is_ko)

    # --- Clause details ---
    heading(L["clauses"])
    if not result.ai_clauses:
        body(L["none"])
    for clause in result.ai_clauses:
        sub = document.add_paragraph()
        r = sub.add_run(f"[{clause.id}] {clause.title}  —  {_v(clause.flag)}")
        _font(r, is_ko, size=12, bold=True)
        kv(L["issue_type"], _v(clause.issue_type))
        kv(L["flag"], _v(clause.flag))
        kv(L["enforceability"], _v(clause.enforceability))
        kv(L["verbatim"], "")
        quote = document.add_paragraph()
        quote.paragraph_format.left_indent = Pt(18)
        _font(quote.add_run(clause.verbatim_text), is_ko, italic=True)
        kv(L["summary"], clause.summary)
        kv(L["impact"], clause.impact_on_you)
        kv(L["revision"], clause.proposed_revision)
        kv(L["legal_basis"], clause.legal_basis)
        document.add_paragraph()

    # --- Benchmarks ---
    heading(L["benchmarks"])
    if result.benchmarks:
        _table(
            document, is_ko,
            [L["topic"], L["contract_value"], L["standard"], L["assessment"]],
            [
                [b.topic, b.contract_value, b.industry_standard, _v(b.assessment)]
                for b in result.benchmarks
            ],
        )
    else:
        body(L["none"])

    # --- Jurisdiction notes ---
    heading(L["jurisdiction_notes"])
    if result.jurisdiction_notes:
        for note in result.jurisdiction_notes:
            sub = document.add_paragraph()
            _font(sub.add_run(note.topic), is_ko, bold=True)
            kv(L["applicable_law"], note.applicable_law)
            kv(L["analysis"], note.analysis)
            kv(L["conflict"], L["yes"] if note.conflict else L["no"])
            document.add_paragraph()
    else:
        body(L["none"])

    # --- Negotiation playbook ---
    heading(L["playbook"])
    if result.negotiation_playbook:
        for item in result.negotiation_playbook:
            tag = "★ walk-away  " if item.walk_away else ""
            sub = document.add_paragraph()
            _font(
                sub.add_run(f"{tag}[{_v(item.priority)}] {item.issue}"),
                is_ko, bold=True,
            )
            kv(L["strategy"], item.strategy)
            kv(L["revision"], item.proposed_revision)
            if item.clause_id:
                kv(L["clause"], item.clause_id)
            document.add_paragraph()
    else:
        body(L["none"])

    # --- Disclaimer ---
    heading(L["disclaimer"])
    body(result.disclaimer)
    body(L["review_note"])

    document.save(str(path))


def _setup_base_style(document, is_ko: bool) -> None:
    from docx.shared import Pt

    style = document.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)
    if is_ko:
        _set_east_asian_font(style.element, _korean_font())


def _font(run, is_ko: bool, *, size: int | None = None, bold: bool = False,
          italic: bool = False) -> None:
    from docx.shared import Pt

    run.font.name = "Times New Roman"
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if is_ko:
        _set_east_asian_font(run._element, _korean_font())


def _set_east_asian_font(element, font_name: str) -> None:
    """Set the w:eastAsia font so Korean glyphs render correctly."""
    from docx.oxml.ns import qn

    rpr = element.get_or_add_rPr() if hasattr(element, "get_or_add_rPr") else element.find(qn("w:rPr"))
    if rpr is None:
        return
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font_name)


def _table(document, is_ko: bool, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, head in enumerate(headers):
        cell = table.rows[0].cells[i]
        _font(cell.paragraphs[0].add_run(head), is_ko, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            _font(cells[i].paragraphs[0].add_run(str(value)), is_ko)


# --------------------------------------------------------------------------- #
# Optional PDF
# --------------------------------------------------------------------------- #
def _convert_to_pdf(docx_path: Path, pdf_path: Path) -> bool:
    """Best-effort DOCX→PDF. Requires Word (docx2pdf). Returns success flag."""
    try:
        from docx2pdf import convert
    except ImportError:
        return False
    try:
        convert(str(docx_path), str(pdf_path))
        return pdf_path.exists()
    except Exception:  # pragma: no cover - environment dependent
        return False
